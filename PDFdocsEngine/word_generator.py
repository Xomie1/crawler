"""
Word document generator using python-docx.
Supports formatting from Google Docs.
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from typing import Optional, List, Dict, Any
import os
import re


class WordGenerator:
    """Generates Word (.docx) documents from text content."""
    
    def __init__(self, output_dir: str = "output"):
        """Initialize Word generator.
        
        Args:
            output_dir: Directory to save generated Word documents
        """
        self.output_dir = output_dir
        os.makedirs(output_dir, exist_ok=True)
    
    def _sanitize_text(self, text: str) -> str:
        """Remove XML-incompatible characters from text.
        
        Removes NULL bytes and control characters that cause issues with python-docx.
        
        Args:
            text: Text to sanitize
            
        Returns:
            Sanitized text safe for XML/Word documents
        """
        if not text:
            return ""
        
        # Remove NULL bytes
        text = text.replace('\x00', '')
        
        # Remove other control characters except newline, tab, and carriage return
        # Keep: \n (0x0A), \r (0x0D), \t (0x09)
        # Remove: other control characters (0x00-0x08, 0x0B-0x0C, 0x0E-0x1F)
        text = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F]', '', text)
        
        # Remove characters that are not valid in XML 1.0
        # XML 1.0 allows: #x9 | #xA | #xD | [#x20-#xD7FF] | [#xE000-#xFFFD] | [#x10000-#x10FFFF]
        # We'll be more permissive and just remove clearly problematic ones
        text = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F-\x84\x86-\x9F]', '', text)
        
        return text
    
    def _set_japanese_font(self, run, font_name: str):
        """Safely set Japanese font for a text run.
        
        Args:
            run: The text run to set font for
            font_name: The font name to use for Japanese characters
        """
        try:
            # Ensure rPr exists
            if run._element.rPr is None:
                from docx.oxml import parse_xml
                run._element.rPr = parse_xml(r'<w:rPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
            
            # Ensure rFonts exists
            if run._element.rPr.rFonts is None:
                from docx.oxml import parse_xml
                run._element.rPr.rFonts = parse_xml(
                    r'<w:rFonts xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
                )
            
            # Set the Japanese font
            run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        except Exception:
            # If setting Japanese font fails, just continue without it
            # The document will still work, just may not display Japanese correctly
            pass
    
    def generate(self, content: Any, filename: str) -> str:
        """Generate Word document from content.
        
        Args:
            content: Either plain text string or list of structured paragraphs from Google Docs
            filename: Output filename (without extension)
            
        Returns:
            Path to generated Word file
        """
        output_path = os.path.join(self.output_dir, f"{filename}.docx")
        doc = Document()
        
        # Set document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Check if content is structured (list of dicts) or plain text
        if isinstance(content, list) and len(content) > 0 and isinstance(content[0], dict):
            # Structured content from Google Docs
            for para_info in content:
                self._add_paragraph_from_formatting(doc, para_info)
        else:
            # Plain text content (backward compatibility)
            paragraphs = self._parse_content(str(content))
            
            for para_text in paragraphs:
                para_text = para_text.strip()
                if not para_text:
                    continue
                
                # Sanitize text to remove XML-incompatible characters
                para_text = self._sanitize_text(para_text)
                if not para_text:
                    continue
                
                # Title
                if para_text.startswith('PRENUPTIAL AGREEMENT') or para_text.startswith('DIVORCE SETTLEMENT AGREEMENT'):
                    p = doc.add_heading(para_text, level=1)
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    self._format_paragraph(p, bold=True, size=16)
                # Headings (numbered sections)
                elif self._is_heading(para_text):
                    p = doc.add_heading(para_text, level=2)
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    self._format_paragraph(p, bold=True, size=12)
                # Body text
                else:
                    p = doc.add_paragraph(para_text)
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    self._format_paragraph(p, bold=False, size=10)
        
        doc.save(output_path)
        return output_path
    
    def _add_paragraph_from_formatting(self, doc: Document, para_info: Dict[str, Any]):
        """Add a paragraph to the document with formatting from Google Docs."""
        text = para_info.get('text', '').strip()
        if not text:
            return
        
        # Get formatting properties
        font_size = para_info.get('font_size') or 11.0  # Ensure not None
        font_family = para_info.get('font_family') or 'Calibri'  # Ensure not None
        bold = para_info.get('bold', False)
        italic = para_info.get('italic', False)
        underline = para_info.get('underline', False)
        alignment_str = para_info.get('alignment', 'LEFT')
        heading = para_info.get('heading')
        text_runs = para_info.get('text_runs', [])
        
        # Convert alignment
        alignment_map = {
            'LEFT': WD_ALIGN_PARAGRAPH.LEFT,
            'CENTER': WD_ALIGN_PARAGRAPH.CENTER,
            'RIGHT': WD_ALIGN_PARAGRAPH.RIGHT,
            'JUSTIFY': WD_ALIGN_PARAGRAPH.JUSTIFY
        }
        alignment = alignment_map.get(alignment_str, WD_ALIGN_PARAGRAPH.LEFT)
        
        # Sanitize text
        text = self._sanitize_text(text)
        if not text:
            return
        
        # Determine if it's a heading
        if heading == 'HEADING_1' or heading == 'TITLE':
            p = doc.add_heading(text, level=1)
        elif heading and heading.startswith('HEADING_'):
            level = int(heading.split('_')[1]) if heading.split('_')[1].isdigit() else 2
            p = doc.add_heading(text, level=min(level, 9))
        else:
            p = doc.add_paragraph()
        
        # Set alignment
        p.alignment = alignment
        
        # Add text with formatting
        if text_runs and len(text_runs) > 1:
            # Multiple text runs with different formatting
            p.clear()  # Clear default run
            for run_info in text_runs:
                run_text = self._sanitize_text(run_info.get('text', ''))
                if not run_text:
                    continue
                
                run = p.add_run(run_text)
                run_font_family = run_info.get('font_family') or font_family
                run_font_size = run_info.get('font_size') or font_size
                run.font.name = run_font_family
                run.font.size = Pt(run_font_size)
                run.font.bold = run_info.get('bold', False)
                run.font.italic = run_info.get('italic', False)
                run.font.underline = run_info.get('underline', False)
                # Set Japanese font
                self._set_japanese_font(run, run.font.name)
        else:
            # Single formatting or no runs
            if not text_runs:
                # Clear default run and add formatted text
                p.clear()
                run = p.add_run(text)
                run.font.name = font_family
                run.font.size = Pt(font_size)
                run.font.bold = bold
                run.font.italic = italic
                run.font.underline = underline
                # Set Japanese font
                self._set_japanese_font(run, font_family)
            else:
                # Use first text run formatting
                run_info = text_runs[0]
                p.clear()
                run = p.add_run(text)
                run_font_family = run_info.get('font_family') or font_family
                run_font_size = run_info.get('font_size') or font_size
                run.font.name = run_font_family
                run.font.size = Pt(run_font_size)
                run.font.bold = run_info.get('bold', bold)
                run.font.italic = run_info.get('italic', italic)
                run.font.underline = run_info.get('underline', underline)
                # Set Japanese font
                self._set_japanese_font(run, run.font.name)
        
        # Set paragraph spacing
        space_before = para_info.get('space_before', 0.0)
        space_after = para_info.get('space_after', 6.0)
        if space_before > 0:
            p.paragraph_format.space_before = Pt(space_before)
        if space_after > 0:
            p.paragraph_format.space_after = Pt(space_after)
        
        # Set line spacing
        line_spacing = para_info.get('line_spacing', 1.15)
        if line_spacing:
            p.paragraph_format.line_spacing = line_spacing
    
    def _is_heading(self, text: str) -> bool:
        """Check if text is a heading (numbered section)."""
        text = text.strip()
        if text and (text[0].isdigit() or text.startswith('WHEREAS') or text.startswith('NOW, THEREFORE')):
            if len(text) < 100 and ('.' in text[:5] or text.startswith('WHEREAS') or text.startswith('NOW, THEREFORE')):
                return True
        return False
    
    def _format_paragraph(self, paragraph, bold: bool = False, size: int = 10):
        """Format paragraph with font settings."""
        for run in paragraph.runs:
            run.font.name = 'Calibri'
            run.font.size = Pt(size)
            run.font.bold = bold
            self._set_japanese_font(run, 'Calibri')
    
    def _parse_content(self, content: str) -> list:
        """Parse content into paragraphs."""
        sections = content.split('\n\n')
        paragraphs = []
        
        for section in sections:
            section = section.strip()
            if not section:
                continue
            
            lines = section.split('\n')
            current_para = []
            
            for line in lines:
                line = line.strip()
                if not line:
                    if current_para:
                        paragraphs.append(' '.join(current_para))
                        current_para = []
                    continue
                
                if self._is_heading(line) and current_para:
                    paragraphs.append(' '.join(current_para))
                    current_para = [line]
                else:
                    current_para.append(line)
            
            if current_para:
                paragraphs.append(' '.join(current_para))
        
        return paragraphs

